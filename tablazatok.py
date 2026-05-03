"""
Ez a fájl a szakdolgozatba illeszthető Word-táblázatok előállításához
szükséges segédfüggvényeket tartalmazza.
A modul szimulációs paraméter-, hozamfelbontási, hasznossági, modell-,
extrémpélda- és rho-komponenstáblázatokat készít docx formátumban.
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT

from benchmark_setup import benchmark_kezdeti
from Black_Scholes import black_scholes_t0, black_scholes_t1
from piaci_allapotvaltozok import benchmark_piaci_allapotvaltozokkal
from hozamok import teljesitmeny_es_hozam_felbontas
from aktiv_portfoliok import aktivan_kezelt_portfoliok
from Brinson import (
    pm,
    modell_eredmenyek,
    portfoliok_hasznossaga,
    szektorsuly_es_aktiv_opciosuly_tablazat_input,
    word_tablak_input,
    pozitiv_allokacio_extrem_peldak_input,
    pozitiv_szelekcio_extrem_peldak_input,
    valtozo_allokacio_extrem_peldak_input,
    rho_adatok
)

mappa_utvonal = r""

def jeloles_beszurasa(cella, reszek, betumeret=11):
    bekezdes = cella.paragraphs[0]
    bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for szoveg, also_index in reszek:
        futas = bekezdes.add_run(szoveg)
        futas.font.size = Pt(betumeret)
        futas.font.subscript = also_index


def szoveg_beszurasa(cella, szoveg, igazitas="left", betumeret=11, felkover=False):
    bekezdes = cella.paragraphs[0]

    if igazitas == "center":
        bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif igazitas == "right":
        bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    futas = bekezdes.add_run(str(szoveg))
    futas.font.size = Pt(betumeret)
    futas.bold = felkover


def tablazat_formazasa(tablazat):
    tablazat.alignment = WD_TABLE_ALIGNMENT.CENTER
    tablazat.style = "Table Grid"

    for sor in tablazat.rows:
        for cella in sor.cells:
            cella.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def szimulacios_parametertablazat_keszitese(mappa_utvonal: str, fajlnev: str):
    teljes_eleresi_ut = os.path.join(mappa_utvonal, fajlnev)

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.top_margin = Cm(2.5)
    szakasz.bottom_margin = Cm(2.5)
    szakasz.left_margin = Cm(2.5)
    szakasz.right_margin = Cm(2.5)

    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run("A szimuláció paraméterei")
    cim_futas.bold = True
    cim_futas.font.size = Pt(13)

    dokumentum.add_paragraph("")

    tablazat = dokumentum.add_table(rows=1, cols=3)
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells
    szoveg_beszurasa(fejlec[0], "Jelölés", igazitas="center", felkover=True)
    szoveg_beszurasa(fejlec[1], "Érték", igazitas="center", felkover=True)
    szoveg_beszurasa(fejlec[2], "Magyarázat", igazitas="center", felkover=True)

    adatok = [
        ([("Véletlen mag", False)], "80", "A reprodukálhatóság biztosítására"),
        ([("N", False)], "11", "A benchmark portfólióban szereplő szektorok száma"),
        ([("M", False)], "10", "Az egy szektoron belül szereplő opciók száma"),
        ([("α", False), ("i", True)], "1.0", "A Dirichlet-eloszlás paramétere a szektorsúlyokhoz"),
        ([("σ", False), ("w", True)], "0.35", "A szektoron belüli opciósúlyok szórása"),
        ([("ξ", False), ("C", True)], "0.5", "A call opciók előfordulási valószínűsége"),
        ([("S", False), ("min", True)], "80", "A kezdeti részvényárfolyam alsó határa"),
        ([("S", False), ("max", True)], "120", "A kezdeti részvényárfolyam felső határa"),
        ([("m", False), ("min", True)], "0.9", "A moneyness intervallum alsó határa"),
        ([("m", False), ("max", True)], "1.1", "A moneyness intervallum felső határa"),
        ([("τ", False), ("min", True)], "30 nap", "A hátralévő futamidő alsó határa"),
        ([("τ", False), ("max", True)], "365 nap", "A hátralévő futamidő felső határa"),
        ([("σ", False), ("min", True)], "0.15", "A kezdeti implikált volatilitás alsó határa"),
        ([("σ", False), ("max", True)], "0.30", "A kezdeti implikált volatilitás felső határa"),
        ([("r", False), ("f0", True)], "0.03", "A kezdeti kockázatmentes kamatláb"),
        ([("Δ", False), ("t", True)], "1/52", "A vizsgált periódus hossza"),
        ([("μ", False), ("min", True)], "-0.05", "A szektor szintű drift alsó határa"),
        ([("μ", False), ("max", True)], "0.10", "A szektor szintű drift felső határa"),
        ([("σ", False), ("μ", True)], "0.03", "Az egyedi drift szórása"),
        ([("ρ", False), ("S", True)], "0.4", "A részvényár-sokkok szektoron belüli korrelációja"),
        ([("ρ", False), ("σ", True)], "0.3", "A volatilitási sokkok szektoron belüli korrelációja"),
        ([("ρ", False), ("SV", True)], "-0.2", "A részvényár és volatilitás kapcsolata"),
        ([("η", False)], "0.25", "A volatilitásdinamika intenzitása"),
        ([("κ", False), ("r", True)], "1.50", "A kamatláb visszahúzó ereje"),
        ([("θ", False), ("r", True)], "0.04", "A kamatláb hosszú távú szintje"),
        ([("σ", False), ("r", True)], "0.01", "A kamatsokk szórása"),
        ([("PM", False)], "100", "A szimulált portfóliómenedzserek száma"),
        ([("λ", False)], "100.0", "A szektorsúlyok koncentrációs paramétere"),
        ([("σ", False), ("aktiv", True)], "0.08", "Az aktív eltérések szórása"),
    ]

    for jeloles, ertek, magyarazat in adatok:
        uj_sor = tablazat.add_row().cells
        jeloles_beszurasa(uj_sor[0], jeloles, betumeret=11)
        szoveg_beszurasa(uj_sor[1], ertek, igazitas="center", betumeret=11)
        szoveg_beszurasa(uj_sor[2], magyarazat, igazitas="left", betumeret=11)

    for sor in tablazat.rows:
        sor.cells[0].width = Cm(3.2)
        sor.cells[1].width = Cm(3.0)
        sor.cells[2].width = Cm(10.3)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def szektor_gorog_tablazat_keszitese(df, mappa_utvonal: str, fajlnev: str):
    teljes_eleresi_ut = os.path.join(mappa_utvonal, fajlnev)

    dokumentum = Document()

    tablazat = dokumentum.add_table(rows=1, cols=8)
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells

    szoveg_beszurasa(fejlec[0], "Név", igazitas="center", felkover=True)
    szoveg_beszurasa(fejlec[1], "Hozam", igazitas="center", felkover=True)

    jeloles_beszurasa(fejlec[2], [("r", False), ("Δ", True)])
    jeloles_beszurasa(fejlec[3], [("r", False), ("Γ", True)])
    jeloles_beszurasa(fejlec[4], [("r", False), ("Θ", True)])
    jeloles_beszurasa(fejlec[5], [("r", False), ("ν", True)])
    jeloles_beszurasa(fejlec[6], [("r", False), ("ρ", True)])
    jeloles_beszurasa(fejlec[7], [("r", False), ("ε", True)])

    szurt_df = df[
        df["nev"].str.startswith("Szektor") | (df["nev"] == "Benchmark")
    ].copy()

    szektorok = szurt_df[szurt_df["nev"].str.startswith("Szektor")].copy()
    szektorok["szektor_id"] = (
        szektorok["nev"].str.extract(r"Szektor (\d+)")[0].astype(int)
    )
    szektorok = szektorok.sort_values("szektor_id")

    benchmark = szurt_df[szurt_df["nev"] == "Benchmark"]

    vegso_df = pd.concat([szektorok, benchmark], ignore_index=True)

    for _, sor in vegso_df.iterrows():
        uj_sor = tablazat.add_row().cells

        szoveg_beszurasa(uj_sor[0], sor["nev"], igazitas="left")
        szoveg_beszurasa(uj_sor[1], f"{float(sor['hozam']):.4f}")

        szoveg_beszurasa(uj_sor[2], f"{float(sor['r_delta']):.4f}")
        szoveg_beszurasa(uj_sor[3], f"{float(sor['r_gamma']):.4f}")
        szoveg_beszurasa(uj_sor[4], f"{float(sor['r_theta']):.4f}")
        szoveg_beszurasa(uj_sor[5], f"{float(sor['r_vega']):.4f}")
        szoveg_beszurasa(uj_sor[6], f"{float(sor['r_rho']):.4f}")
        szoveg_beszurasa(uj_sor[7], f"{float(sor['r_epsilon']):.4f}")

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def hasznossagi_kategoriak(lokacio: str, fajlnev: str) -> None:
    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.top_margin = Cm(2.5)
    szakasz.bottom_margin = Cm(2.5)
    szakasz.left_margin = Cm(2.5)
    szakasz.right_margin = Cm(2.5)

    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run("Hasznossági kategóriák")
    cim_futas.bold = True
    cim_futas.font.size = Pt(13)

    dokumentum.add_paragraph("")

    tablazat = dokumentum.add_table(rows=1, cols=2)
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells
    szoveg_beszurasa(fejlec[0], "Kategória", igazitas="center", betumeret=11, felkover=True)
    szoveg_beszurasa(fejlec[1], "Feltétel", igazitas="center", betumeret=11, felkover=True)

    adatok = [
        ("Nem hasznos", "H_pm < 0.2"),
        ("Gyengén hasznos", "0.2 ≤ H_pm < 0.4"),
        ("Közepesen hasznos", "0.4 ≤ H_pm < 0.6"),
        ("Hasznos", "0.6 ≤ H_pm < 0.8"),
        ("Nagyon hasznos", "0.8 ≤ H_pm ≤ 1"),
    ]

    for kategoria, feltetel in adatok:
        uj_sor = tablazat.add_row().cells
        szoveg_beszurasa(uj_sor[0], kategoria, igazitas="left", betumeret=11)
        szoveg_beszurasa(uj_sor[1], feltetel, igazitas="center", betumeret=11)

    for sor in tablazat.rows:
        sor.cells[0].width = Cm(5.0)
        sor.cells[1].width = Cm(11.5)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def hasznossagi_kategoriak_gyakorisaga_tablazat(
    eredmenyek,
    mappa_utvonal: str,
    fajlnev: str,
    modell_index: int = 3,
    kuszob: float = 0.3,
    alfa: float = 0.7
):
    teljes_eleresi_ut = os.path.join(mappa_utvonal, fajlnev)

    _, osszesites_df = portfoliok_hasznossaga(
        eredmenyek=eredmenyek,
        modell_index=modell_index,
        kuszob=kuszob,
        alfa=alfa
    )

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.top_margin = Cm(2.5)
    szakasz.bottom_margin = Cm(2.5)
    szakasz.left_margin = Cm(2.5)
    szakasz.right_margin = Cm(2.5)

    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run("A portfóliók hasznossági kategóriáinak megoszlása")
    cim_futas.bold = True
    cim_futas.font.size = Pt(13)

    dokumentum.add_paragraph("")

    tablazat = dokumentum.add_table(rows=1, cols=2)
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells
    szoveg_beszurasa(fejlec[0], "Kategória", igazitas="center", betumeret=11, felkover=True)
    szoveg_beszurasa(fejlec[1], "Darabszám", igazitas="center", betumeret=11, felkover=True)

    for _, sor in osszesites_df.iterrows():
        uj_sor = tablazat.add_row().cells
        szoveg_beszurasa(uj_sor[0], sor["Kategória"], igazitas="left", betumeret=11)
        szoveg_beszurasa(uj_sor[1], int(sor["Darabszám"]), igazitas="center", betumeret=11)

    for sor in tablazat.rows:
        sor.cells[0].width = Cm(9.0)
        sor.cells[1].width = Cm(4.0)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def osszetett_jeloles_beszurasa(cella, alap: str, felso: str = "", also: str = "", betumeret: int = 11):
    bekezdes = cella.paragraphs[0]
    bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    futas = bekezdes.add_run(alap)
    futas.font.size = Pt(betumeret)

    if felso:
        futas = bekezdes.add_run(felso)
        futas.font.size = Pt(betumeret)
        futas.font.superscript = True

    if also:
        futas = bekezdes.add_run(also)
        futas.font.size = Pt(betumeret)
        futas.font.subscript = True


def aktiv_gorog_felbontas_sablon(lokacio: str, fajlnev: str) -> None:
    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.top_margin = Cm(2.5)
    szakasz.bottom_margin = Cm(2.5)
    szakasz.left_margin = Cm(2.5)
    szakasz.right_margin = Cm(2.5)

    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run("Aktív görög betűk szerinti felbontás")
    cim_futas.bold = True
    cim_futas.font.size = Pt(13)

    dokumentum.add_paragraph("")

    tablazat = dokumentum.add_table(rows=7, cols=13)
    tablazat_formazasa(tablazat)

    elso_sor = tablazat.rows[0].cells
    masodik_sor = tablazat.rows[1].cells

    szoveg_beszurasa(elso_sor[0], "", igazitas="center", betumeret=11, felkover=True)

    allokacio_cella = elso_sor[1]
    for i in range(2, 7):
        allokacio_cella = allokacio_cella.merge(elso_sor[i])
    szoveg_beszurasa(allokacio_cella, "Allokáció", igazitas="center", betumeret=11, felkover=True)

    szelekcio_cella = elso_sor[7]
    for i in range(8, 13):
        szelekcio_cella = szelekcio_cella.merge(elso_sor[i])
    szoveg_beszurasa(szelekcio_cella, "Szelekció", igazitas="center", betumeret=11, felkover=True)

    szoveg_beszurasa(masodik_sor[0], "", igazitas="center", betumeret=11, felkover=True)

    also_indexek = ["d", "g", "t", "v", "r", "e"]

    for i, idx in enumerate(also_indexek, start=1):
        osszetett_jeloles_beszurasa(masodik_sor[i], "A", "A", idx, 11)

    for i, idx in enumerate(also_indexek, start=7):
        osszetett_jeloles_beszurasa(masodik_sor[i], "A", "S", idx, 11)

    szektor_nevek = ["Szektor 1", "...", "Szektor i", "...", "Szektor N"]

    for sor_index in range(2, 7):
        sor = tablazat.rows[sor_index].cells
        szoveg_beszurasa(sor[0], szektor_nevek[sor_index - 2], igazitas="left", betumeret=11)

        for oszlop_index in range(1, 13):
            szoveg_beszurasa(sor[oszlop_index], "", igazitas="center", betumeret=11)

    for sor in tablazat.rows:
        sor.cells[0].width = Cm(2.8)
        for i in range(1, 13):
            sor.cells[i].width = Cm(1.35)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def szektorsuly_es_aktiv_opciosuly_tablazat(
    benchmark_df: pd.DataFrame,
    aktiv_portfolio_df: pd.DataFrame,
    lokacio: str,
    fajlnev: str
) -> None:
    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    df = szektorsuly_es_aktiv_opciosuly_tablazat_input(
        benchmark_df=benchmark_df,
        aktiv_portfolio_df=aktiv_portfolio_df
    )

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.top_margin = Cm(2.5)
    szakasz.bottom_margin = Cm(2.5)
    szakasz.left_margin = Cm(2.5)
    szakasz.right_margin = Cm(2.5)

    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run("Szektorsúlyok és szektoron belüli aktív opciósúlyok")
    cim_futas.bold = True
    cim_futas.font.size = Pt(13)

    dokumentum.add_paragraph("")

    tablazat = dokumentum.add_table(rows=1, cols=5)
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells
    szoveg_beszurasa(fejlec[0], "Szektor", igazitas="center", betumeret=11, felkover=True)
    szoveg_beszurasa(fejlec[1], "Portfólió", igazitas="center", betumeret=11, felkover=True)
    szoveg_beszurasa(fejlec[2], "Benchmark", igazitas="center", betumeret=11, felkover=True)
    szoveg_beszurasa(fejlec[3], "Aktív", igazitas="center", betumeret=11, felkover=True)
    szoveg_beszurasa(fejlec[4], "Szektoron belüli átlagos abszolút aktív opciósúly", igazitas="center", betumeret=11, felkover=True)

    for _, sor in df.iterrows():
        uj_sor = tablazat.add_row().cells

        szoveg_beszurasa(uj_sor[0], sor["Szektor"], igazitas="left", betumeret=11)
        szoveg_beszurasa(uj_sor[1], f"{float(sor['Portfólió']):.4f}", igazitas="center", betumeret=11)
        szoveg_beszurasa(uj_sor[2], f"{float(sor['Benchmark']):.4f}", igazitas="center", betumeret=11)
        szoveg_beszurasa(uj_sor[3], f"{float(sor['Aktív']):.4f}", igazitas="center", betumeret=11)
        szoveg_beszurasa(
            uj_sor[4],
            f"{float(sor['Szektoron belüli átlagos aktív opciósúly']):.4f}",
            igazitas="center",
            betumeret=11
        )

    for sor in tablazat.rows:
        sor.cells[0].width = Cm(3.0)
        sor.cells[1].width = Cm(3.0)
        sor.cells[2].width = Cm(3.0)
        sor.cells[3].width = Cm(3.0)
        sor.cells[4].width = Cm(6.0)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def word_multiindex_tablazat_beszurasa(dokumentum, df: pd.DataFrame, cim: str) -> None:
    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run(cim)
    cim_futas.bold = True
    cim_futas.font.size = Pt(12)

    dokumentum.add_paragraph("")

    oszlopok_szama = len(df.columns)
    tablazat = dokumentum.add_table(rows=2, cols=oszlopok_szama)
    tablazat_formazasa(tablazat)

    felso = tablazat.rows[0].cells
    also = tablazat.rows[1].cells

    j = 0
    while j < oszlopok_szama:
        felso_szint = str(df.columns[j][0])
        also_szint = str(df.columns[j][1])

        start = j

        while j + 1 < oszlopok_szama and str(df.columns[j + 1][0]) == felso_szint:
            j += 1

        if j > start:
            cella = felso[start].merge(felso[j])
            cella.text = ""
            szoveg_beszurasa(cella, felso_szint, igazitas="center", betumeret=8, felkover=True)
        else:
            if also_szint == "":
                cella = felso[start].merge(also[start])
                cella.text = ""
                szoveg_beszurasa(cella, felso_szint, igazitas="center", betumeret=8, felkover=True)
            else:
                felso[start].text = ""
                szoveg_beszurasa(felso[start], felso_szint, igazitas="center", betumeret=8, felkover=True)

        j += 1

    for j, col in enumerate(df.columns):
        also_szint = str(col[1])

        if also_szint != "":
            also[j].text = ""
            szoveg_beszurasa(also[j], also_szint, igazitas="center", betumeret=8, felkover=True)

    for _, sor in df.iterrows():
        uj_sor = tablazat.add_row().cells

        for j, col in enumerate(df.columns):
            ertek = sor[col]

            if isinstance(ertek, float):
                szoveg = f"{ertek:.4f}"
            else:
                szoveg = str(ertek)

            igazitas = "left" if j == 0 else "center"
            szoveg_beszurasa(uj_sor[j], szoveg, igazitas=igazitas, betumeret=8)

    for sor in tablazat.rows:
        for j in range(oszlopok_szama):
            if j == 0:
                sor.cells[j].width = Cm(2.7)
            else:
                sor.cells[j].width = Cm(1.8)

    dokumentum.add_page_break()


def pm_modell_tablazatok_word(
    eredmenyek,
    pm_sorszam: int,
    lokacio: str,
    fajlnev: str
) -> None:
    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    tablazatok = word_tablak_input(
        eredmenyek=eredmenyek,
        pm_sorszam=pm_sorszam
    )

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.orientation = WD_ORIENT.LANDSCAPE
    szakasz.page_width, szakasz.page_height = szakasz.page_height, szakasz.page_width
    szakasz.top_margin = Cm(1.5)
    szakasz.bottom_margin = Cm(1.5)
    szakasz.left_margin = Cm(1.5)
    szakasz.right_margin = Cm(1.5)

    for cim, df in tablazatok.items():
        word_multiindex_tablazat_beszurasa(
            dokumentum=dokumentum,
            df=df,
            cim=cim
        )

    if dokumentum.paragraphs[-1]._element.getprevious() is not None:
        utolso = dokumentum.paragraphs[-1]._element
        utolso.getparent().remove(utolso)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def komponens_extrem_peldak_tablazat(
    eredmenyek,
    lokacio: str,
    fajlnev: str
) -> None:
    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    pozitiv_allokacio_df = pozitiv_allokacio_extrem_peldak_input(
        eredmenyek=eredmenyek
    )

    pozitiv_szelekcio_df = pozitiv_szelekcio_extrem_peldak_input(
        eredmenyek=eredmenyek
    )

    valtozo_allokacio_df = valtozo_allokacio_extrem_peldak_input(
        eredmenyek=eredmenyek
    )

    dokumentum = Document()

    szakasz = dokumentum.sections[0]
    szakasz.orientation = WD_ORIENT.LANDSCAPE
    szakasz.page_width, szakasz.page_height = szakasz.page_height, szakasz.page_width
    szakasz.top_margin = Cm(1.5)
    szakasz.bottom_margin = Cm(1.5)
    szakasz.left_margin = Cm(1.5)
    szakasz.right_margin = Cm(1.5)

    komponens_extrem_egy_tablazat_beszurasa(
        dokumentum=dokumentum,
        df=pozitiv_allokacio_df,
        cim="Pozitív többlet felbontás – allokációs szélső példák"
    )

    dokumentum.add_paragraph("")

    komponens_extrem_egy_tablazat_beszurasa(
        dokumentum=dokumentum,
        df=pozitiv_szelekcio_df,
        cim="Pozitív többlet felbontás – szelekciós szélső példák"
    )

    dokumentum.add_paragraph("")

    komponens_extrem_egy_tablazat_beszurasa(
        dokumentum=dokumentum,
        df=valtozo_allokacio_df,
        cim="Változó többlet felbontás – allokációs szélső példák"
    )

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def komponens_extrem_egy_tablazat_beszurasa(
    dokumentum,
    df: pd.DataFrame,
    cim: str
) -> None:
    cim_bekezdes = dokumentum.add_paragraph()
    cim_bekezdes.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cim_futas = cim_bekezdes.add_run(cim)
    cim_futas.bold = True
    cim_futas.font.size = Pt(12)

    dokumentum.add_paragraph("")

    tablazat = dokumentum.add_table(rows=1, cols=len(df.columns))
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells

    for i, oszlop in enumerate(df.columns):
        szoveg_beszurasa(
            fejlec[i],
            oszlop,
            igazitas="center",
            betumeret=8,
            felkover=True
        )

    for _, sor in df.iterrows():
        uj_sor = tablazat.add_row().cells

        for i, oszlop in enumerate(df.columns):
            ertek = sor[oszlop]

            if oszlop == "PM":
                szoveg = str(int(ertek))
            elif isinstance(ertek, float):
                szoveg = f"{ertek:.4f}"
            else:
                szoveg = str(ertek)

            igazitas = "left" if oszlop == "Szektor" else "center"

            szoveg_beszurasa(
                uj_sor[i],
                szoveg,
                igazitas=igazitas,
                betumeret=8
            )

    for sor in tablazat.rows:
        for i, oszlop in enumerate(df.columns):
            if oszlop == "PM":
                sor.cells[i].width = Cm(1.0)
            elif oszlop == "Szektor":
                sor.cells[i].width = Cm(2.0)
            elif oszlop in ["Portfóliósúly", "Benchmarksúly"]:
                sor.cells[i].width = Cm(2.0)
            elif "Brinson" in oszlop:
                sor.cells[i].width = Cm(3.4)
            else:
                sor.cells[i].width = Cm(1.6)


def osszefoglalo_aktiv_tablazat(
    pm_id: int,
    hozam_portfolio: float,
    hozam_benchmark: float,
    aktiv_hozam: float,
    allokacio: float,
    szelekcio: float,
    interakcio: float,
    allokacios_komponensek: dict,
    szelekcios_komponensek: dict,
    lokacio: str,
    fajlnev: str
) -> None:

    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    doc = Document()

    tablazat = doc.add_table(rows=6, cols=7)
    tablazat_formazasa(tablazat)

    fejlec = [
        "PM", "Portfólió hozam", "Benchmark hozam",
        "Aktív hozam", "Allokáció", "Szelekció", "Interakció"
    ]

    for j, szoveg in enumerate(fejlec):
        szoveg_beszurasa(tablazat.rows[0].cells[j], szoveg, "center", 10, True)

    ertekek = [
        pm_id,
        hozam_portfolio,
        hozam_benchmark,
        aktiv_hozam,
        allokacio,
        szelekcio,
        interakcio
    ]

    for j, val in enumerate(ertekek):
        if isinstance(val, float):
            szoveg = f"{val:.4f}"
        else:
            szoveg = str(val)
        szoveg_beszurasa(tablazat.rows[1].cells[j], szoveg, "center", 10)

    fejlec_allok = [
        "Allokáció", "Aktív delta", "Aktív gamma",
        "Aktív theta", "Aktív vega", "Aktív rho", "Aktív residual"
    ]

    for j, szoveg in enumerate(fejlec_allok):
        szoveg_beszurasa(tablazat.rows[2].cells[j], szoveg, "center", 10, True)

    ertekek_allok = [
        "-",
        allokacios_komponensek["delta"],
        allokacios_komponensek["gamma"],
        allokacios_komponensek["theta"],
        allokacios_komponensek["vega"],
        allokacios_komponensek["rho"],
        allokacios_komponensek["residual"],
    ]

    for j, val in enumerate(ertekek_allok):
        if isinstance(val, float):
            szoveg = f"{val:.4f}"
        else:
            szoveg = str(val)
        szoveg_beszurasa(tablazat.rows[3].cells[j], szoveg, "center", 10)

    fejlec_szel = [
        "Szelekció", "Aktív delta", "Aktív gamma",
        "Aktív theta", "Aktív vega", "Aktív rho", "Aktív residual"
    ]

    for j, szoveg in enumerate(fejlec_szel):
        szoveg_beszurasa(tablazat.rows[4].cells[j], szoveg, "center", 10, True)

    ertekek_szel = [
        "-",
        szelekcios_komponensek["delta"],
        szelekcios_komponensek["gamma"],
        szelekcios_komponensek["theta"],
        szelekcios_komponensek["vega"],
        szelekcios_komponensek["rho"],
        szelekcios_komponensek["residual"],
    ]

    for j, val in enumerate(ertekek_szel):
        if isinstance(val, float):
            szoveg = f"{val:.4f}"
        else:
            szoveg = str(val)
        szoveg_beszurasa(tablazat.rows[5].cells[j], szoveg, "center", 10)

    doc.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def ket_opcios_tablazat(
    lokacio: str,
    fajlnev: str
) -> None:

    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    doc = Document()

    tablazat = doc.add_table(rows=1, cols=4)
    tablazat_formazasa(tablazat)

    fejlec = ["Név", "Portfólió súly", "Benchmark súly", "Opció ár"]

    for j, szoveg in enumerate(fejlec):
        szoveg_beszurasa(tablazat.rows[0].cells[j], szoveg, "center", 11, True)

    adatok = [
        ["Long Call", 1.5, -0.5, 31.5766],
        ["Long Put", -0.5, 1.5, 28.6579],
    ]

    for sor in adatok:
        uj_sor = tablazat.add_row().cells
        for j, val in enumerate(sor):
            if isinstance(val, float):
                szoveg = f"{val:.4f}"
            else:
                szoveg = str(val)
            igazitas = "left" if j == 0 else "center"
            szoveg_beszurasa(uj_sor[j], szoveg, igazitas, 11)

    doc.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def opcios_tablazat_t0_K(lokacio: str, fajlnev: str) -> None:
    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)

    doc = Document()

    tablazat = doc.add_table(rows=1, cols=5)
    tablazat_formazasa(tablazat)

    fejlec = tablazat.rows[0].cells
    szoveg_beszurasa(fejlec[0], "Név", "center", 11, True)
    szoveg_beszurasa(fejlec[1], "Portfólió súly", "center", 11, True)
    szoveg_beszurasa(fejlec[2], "Benchmark súly", "center", 11, True)
    szoveg_beszurasa(fejlec[3], "K", "center", 11, True)
    szoveg_beszurasa(fejlec[4], "Opció ár", "center", 11, True)

    adatok = [
        ["Szektor 1", 0.4000, 0.5000, "", ""],
        ["Call", 1.5000, -0.5000, 90, 14.8065],
        ["Put", -0.5000, 1.5000, 90, 3.0244],
        ["Szektor 2", 0.6000, 0.5000, "", ""],
        ["Call", 1.5000, -0.5000, 80, 22.5429],
        ["Put", -0.5000, 1.5000, 80, 0.9588],
    ]

    for sor in adatok:
        uj_sor = tablazat.add_row().cells
        for j, ertek in enumerate(sor):
            if isinstance(ertek, float):
                szoveg = f"{ertek:.4f}"
            else:
                szoveg = str(ertek)

            felkover = str(sor[0]).startswith("Szektor")
            igazitas = "left" if j == 0 else "center"

            szoveg_beszurasa(uj_sor[j], szoveg, igazitas, 11, felkover)

    for sor in tablazat.rows:
        for j in range(5):
            sor.cells[j].width = Cm(3.0)

    doc.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def rho_osszefoglalo_tablazat(lokacio: str, fajlnev: str):

    df = rho_adatok()

    szukseges = df[[
        "Név",
        "Új opció ár",
        "Portfólió hozam",
        "Benchmark hozam",
        "BF 1985 allokáció",
        "Szelekció",
        "Interakció",
        "BHB 1986 allokáció"
    ]]

    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)
    dokumentum = Document()

    tablazat = dokumentum.add_table(rows=1, cols=len(szukseges.columns))
    tablazat.style = "Table Grid"

    fejlec = tablazat.rows[0].cells
    for i, col in enumerate(szukseges.columns):
        p = fejlec[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.size = Pt(10)

    for _, sor in szukseges.iterrows():
        uj_sor = tablazat.add_row().cells
        for j, col in enumerate(szukseges.columns):
            ertek = sor[col]

            if isinstance(ertek, float):
                szoveg = f"{ertek:.4f}"
            elif ertek is None:
                szoveg = ""
            else:
                szoveg = str(ertek)

            p = uj_sor[j].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if j > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(szoveg)
            run.font.size = Pt(10)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


def rho_komponens_tablazat(lokacio: str, fajlnev: str):
    import os
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    df = rho_adatok()

    df_szurt = df[
        df["Név"].str.startswith("Szektor") | (df["Név"] == "Összesen")
    ]

    szukseges = df_szurt[[
        "Név",
        "Szel. θ",
        "Szel. ρ",
        "Szel. ε",
        "Vált. allok. θ",
        "Vált. allok. ρ",
        "Vált. allok. ε",
        "Poz. allok. θ",
        "Poz. allok. ρ",
        "Poz. allok. ε"
    ]]

    teljes_eleresi_ut = os.path.join(lokacio, fajlnev)
    dokumentum = Document()

    tablazat = dokumentum.add_table(rows=1, cols=len(szukseges.columns))
    tablazat.style = "Table Grid"

    fejlec = tablazat.rows[0].cells
    for i, col in enumerate(szukseges.columns):
        p = fejlec[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.size = Pt(10)

    for _, sor in szukseges.iterrows():
        uj_sor = tablazat.add_row().cells
        for j, col in enumerate(szukseges.columns):
            ertek = sor[col]

            if isinstance(ertek, float):
                szoveg = f"{ertek:.4f}"
            else:
                szoveg = str(ertek)

            p = uj_sor[j].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if j > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(szoveg)
            run.font.size = Pt(10)

    dokumentum.save(teljes_eleresi_ut)
    print(f"Elkészült: {teljes_eleresi_ut}")


if __name__ == "__main__":
    #df = benchmark_kezdeti()
    #df = black_scholes_t0(df)
    #df = benchmark_piaci_allapotvaltozokkal(df)
    #df = black_scholes_t1(df)
    #df = teljesitmeny_es_hozam_felbontas(df)

    #szimulacios_parametertablazat_keszitese(mappa_utvonal, "parametertablazat")
    #szektor_gorog_tablazat_keszitese(df, mappa_utvonal, "benchmark_hozamfelbontas.docx")
    #hasznossagi_kategoriak(mappa_utvonal, "hasznossagi_kategoriak.docx")

    #aktiv_portfoliok = aktivan_kezelt_portfoliok(df, pm)
    #negy_modell_eredmenyei = modell_eredmenyek(df, aktiv_portfoliok)

    """
    hasznossagi_kategoriak_gyakorisaga_tablazat(
        eredmenyek=negy_modell_eredmenyei,
        mappa_utvonal=mappa_utvonal,
        fajlnev="hasznossagi_kategoriak_gyakorisaga.docx",
        modell_index=3,
        kuszob=0.3,
        alfa=0.5
    )
    """

    #aktiv_gorog_felbontas_sablon(mappa_utvonal, "tobbletinformacio.docx")

    """
    szektorsuly_es_aktiv_opciosuly_tablazat(
    benchmark_df=df,
    aktiv_portfolio_df=aktiv_portfoliok[62],
    lokacio=mappa_utvonal,
    fajlnev="szektorsuly_aktiv_opciosuly_PM63.docx"
    )   
    """
    """
    pm_modell_tablazatok_word(
    eredmenyek=negy_modell_eredmenyei,
    pm_sorszam=63,
    lokacio=mappa_utvonal,
    fajlnev="PM63_modell_tablazatok.docx"
    )
    """
    """
    komponens_extrem_peldak_tablazat(
    eredmenyek=negy_modell_eredmenyei,
    lokacio=mappa_utvonal,
    fajlnev="komponens_extrem_peldak_3tabla.docx"
    )
    """
    """
    osszefoglalo_aktiv_tablazat(
        pm_id=63,
        hozam_portfolio=0.029041001,
        hozam_benchmark=0.022698184,
        aktiv_hozam=0.006342817,
        allokacio=0.020538162,
        szelekcio=-0.004939084,
        interakcio=-0.009256261,

        allokacios_komponensek={
            "delta": 0.0117627,
            "gamma": 0.005233897,
            "theta": -0.001630927,
            "vega": 0.005101225,
            "rho": 0.000236628,
            "residual": -0.000165361,
        },

        szelekcios_komponensek={
            "delta": 0.006242823,
            "gamma": -0.003480561,
            "theta": -0.002217427,
            "vega": -0.005653038,
            "rho": 0.000146771,
            "residual": 2.23479e-05,
        },

        lokacio=mappa_utvonal,
        fajlnev="teljes_portfolio.docx"
        )
        """
    
    #opcios_tablazat_t0_K(mappa_utvonal,"rho_alap.docx")
    #rho_osszefoglalo_tablazat(mappa_utvonal,"rho_sima.docx")
    #rho_komponens_tablazat(mappa_utvonal,"rho_felbontas.docx")
    
    